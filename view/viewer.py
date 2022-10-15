import docx
from docx.shared import Inches
import os
from config import report_prefix_path
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.oxml.ns import qn
from common.translate import Translator
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


class ArticleViewer:

    def __init__(self, article):
        """
        需要article作为参数
        :param article: article数据结构
        """
        self.article = article

    def publish_en_report(self, result_path=None):
        # 传值
        aid = self.article.aid
        website = self.article.website
        kind = self.article.kind
        body = self.article.body
        publish_date = self.article.publish_date
        title = self.article.title
        url = self.article.url
        author = self.article.author
        keyword = self.article.keyword
        attachment = self.article.attachment

        special_char = "\\ / : * ? \" \" \' \' < > |".split(" ")
        # 去除title中的特殊字符
        for char in special_char:
            if char in title:
                title = title.replace(char, "")

        if result_path is None:
            result_path_pre = os.path.join(report_prefix_path, "en/{}/{}-{}/".format(publish_date[:7], website, kind))
            doc_name = "{}-{}.docx".format(publish_date.replace(" ", ""), title)
            if not os.path.exists(result_path_pre):
                os.makedirs(result_path_pre)
            result_path = result_path_pre + doc_name

        # 新建文档对象按模板新建 word 文档文件，具有模板文件的所有格式
        doc = docx.Document()

        # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
        # en_style = doc.styles.add_style('en', 1)

        # 增加标题:add_heading(self, text="", level=1):
        doc.add_heading(text=title, level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加作者
        doc.add_paragraph(f'Author:{author}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 处理源码字符串，转换为文本
        if body is not None:
            body = BeautifulSoup(body, "html.parser")

            # 只查找下一级孩子，不需要递归
            if body.div is not None:
                # 针对ecb working paper
                if website == "ECB" and kind == "working_paper":
                    tag_lst = body.div.dl.find_all(recursive=False)[:-2]
                else:
                    tag_lst = body.div.find_all(recursive=False)

            elif body.span is not None:
                tag_lst = body.span.find_all(recursive=False)
            else:
                raise AttributeError(f"{url} 网站html格式不正确")
            for tag in tag_lst:
                if "h" in tag.name:
                    # 加粗
                    para = doc.add_paragraph()
                    para.add_run(tag.text.strip()).bold = True
                # 如果子标签也有div
                # elif "div" in tag.name:
                #     temp_tag_lst = tag.find_all(recursive=False)
                #     for temp_tag in temp_tag_lst:
                #         para = doc.add_paragraph()
                #         if temp_tag.string is not None:
                #             print(temp_tag.string)
                #             para.add_run(temp_tag.string.strip())
                else:
                    # print(tag.text)
                    # 删去figure,table的部分
                    if "Figure" not in tag.text:
                        if tag.get('class') is not None:

                            tag_class = tag.get('class')[0]
                            if "table" in tag_class:
                                continue
                        para = doc.add_paragraph(tag.text.strip())
                        para.paragraph_format.first_line_indent = Pt(10)  # 首行缩进10磅

                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # text_list = body.text.split("\n")
        #
        # for text in text_list:
        #     if not text.isspace():
        #         para = doc.add_paragraph(text=text.strip())
        #         para.paragraph_format.first_line_indent = Pt(10)  # 首行缩进10磅
        #         para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # para = doc.add_paragraph(text=body.text)
        # para.paragraph_format.first_line_indent = Pt(10)  # 首行缩进10磅

        for i in range(1):
            doc.add_paragraph()

        # 添加url
        if url is not None:
            p = doc.add_paragraph()
            p.add_run(f"Url:\n").bold = True
            # 在段落中添加文字块，add_run(self, text=None, style=None):返回一个 run 对象
            hyperlink = add_hyperlink(p, url, url, '0000FF', underline=True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            p = doc.add_paragraph()
            p.add_run("Url:").bold = True
            p.add_run("NA")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 添加附件，如果有就写上
        if attachment is not None:
            p = doc.add_paragraph()
            p.add_run("PDF:").bold = True
            # 在段落中添加文字块，add_run(self, text=None, style=None):返回一个 run 对象
            hyperlink = add_hyperlink(p, attachment, attachment, '0000FF', underline=True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            p = doc.add_paragraph()
            p.add_run("PDF:").bold = True
            p.add_run("NA")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 添加单位
        p = doc.add_paragraph()
        p.add_run(f"From:").bold = True
        p.add_run(f"{website} - {kind}")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 保存文件
        doc.save(result_path)

    def publish_cn_report(self, result_path=None):
        # 传值
        aid = self.article.aid
        website = self.article.website
        kind = self.article.kind
        body = self.article.body
        publish_date = self.article.publish_date
        title = self.article.title
        url = self.article.url
        author = self.article.author
        keyword = self.article.keyword
        attachment = self.article.attachment

        special_char = "\\ / : * ? \" \" \' \' < > |".split(" ")
        # 去除title中的特殊字符
        for char in special_char:
            if char in title:
                title = title.replace(char, "")

        if result_path is None:
            result_path_pre = os.path.join(report_prefix_path,
                                           "cn/{}/{}-{}/".format(publish_date[:7], website, kind))
            doc_name = "{}-{}.docx".format(publish_date.replace(" ", ""), title)
            if not os.path.exists(result_path_pre):
                os.makedirs(result_path_pre)
            result_path = result_path_pre + doc_name

        # 新建文档对象按模板新建 word 文档文件，具有模板文件的所有格式
        doc = docx.Document()

        # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
        cn_style = doc.styles.add_style('cn', 1)
        # 设置中文字体
        cn_style.font.name = '微软雅黑'
        cn_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 增加标题:add_heading(self, text="", level=1):
        doc.add_heading(text=Translator.translate(title.strip()), level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加作者
        doc.add_paragraph(f'作者:{author}', style=cn_style).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 处理源码字符串，转换为文本
        if body is not None:
            body = BeautifulSoup(body, "html.parser")

            # 只查找下一级孩子，不需要递归
            if body.div is not None:
                # 针对ecb working paper
                if website == "ECB" and kind == "working_paper":
                    tag_lst = body.div.dl.find_all(recursive=False)[:-2]
                else:
                    tag_lst = body.div.find_all(recursive=False)

            elif body.span is not None:
                tag_lst = body.span.find_all(recursive=False)
            else:
                raise AttributeError(f"{url} 网站html格式不正确")

            for tag in tag_lst:
                if "h" in tag.name:
                    # 加粗
                    # 保证输入不为空
                    if tag.text != "":
                        para = doc.add_paragraph(style=cn_style)
                        para.add_run(Translator.translate(tag.text)).bold = True
                else:
                    # 删去figure,table的部分
                    if "Figure" not in tag.text:
                        if tag.get('class') is not None:
                            tag_class = tag.get('class')[0]
                            if "table" in tag_class:
                                continue

                        # 保证输入不为空
                        text = tag.text.strip()
                        if text != "":
                            para = doc.add_paragraph(Translator.translate(text), style=cn_style)
                            para.paragraph_format.first_line_indent = Pt(20)  # 首行缩进20磅

                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for i in range(1):
            doc.add_paragraph()

        # 添加url
        if url is not None:
            p = doc.add_paragraph(style=cn_style)
            p.add_run(f"原文链接:\n").bold = True
            # 在段落中添加文字块，add_run(self, text=None, style=None):返回一个 run 对象
            hyperlink = add_hyperlink(p, url, url, '0000FF', underline=True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            p = doc.add_paragraph(style=cn_style)
            p.add_run("原文链接:").bold = True
            p.add_run("无")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # 保存文件
        doc.save(result_path)
